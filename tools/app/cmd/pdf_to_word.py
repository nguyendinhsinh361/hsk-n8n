# Cài đặt thư viện cần thiết:
# pip install pytesseract pdf2image python-docx pillow concurrent-log-handler torch

import os
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import concurrent.futures
import time
import logging
import torch
import io
from PIL import Image, ImageEnhance

# Kiểm tra xem MPS (Metal Performance Shaders) có khả dụng không
mps_available = hasattr(torch.backends, 'mps') and torch.backends.mps.is_available()

# Thiết lập logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def optimize_image(image, quality=85, max_width=1200):
    """
    Tối ưu hóa hình ảnh để giảm kích thước file nhưng vẫn giữ chất lượng tốt
    
    :param image: Đối tượng hình ảnh PIL
    :param quality: Chất lượng nén JPEG (1-100)
    :param max_width: Chiều rộng tối đa của hình ảnh
    :return: Hình ảnh đã được tối ưu hóa
    """
    # Điều chỉnh kích thước nếu hình ảnh quá lớn
    width, height = image.size
    if width > max_width:
        ratio = max_width / width
        new_width = max_width
        new_height = int(height * ratio)
        image = image.resize((new_width, new_height), Image.LANCZOS)
    
    # Tăng độ tương phản và độ sắc nét để hình ảnh rõ ràng hơn
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(1.3)  # Tăng độ tương phản
    
    enhancer = ImageEnhance.Sharpness(image)
    image = enhancer.enhance(1.3)  # Tăng độ sắc nét
    
    # Chuyển đổi sang định dạng JPEG để giảm kích thước
    if image.mode != 'RGB':
        image = image.convert('RGB')
    
    # Lưu hình ảnh vào bộ nhớ tạm để tối ưu hóa
    with io.BytesIO() as output:
        image.save(output, format='JPEG', quality=quality, optimize=True)
        output.seek(0)
        optimized_image = Image.open(output)
        optimized_image.load()  # Đảm bảo hình ảnh được tải vào bộ nhớ
    
    return optimized_image

def process_page(args):
    """
    Hàm xử lý một trang PDF
    
    :param args: Tuple chứa (image, page_num, total_pages)
    :return: Tuple (page_num, text, image, processing_time)
    """
    image, page_num, total_pages = args
    start_time = time.time()
    
    logger.info(f"Đang xử lý trang {page_num+1}/{total_pages}...")
    
    # Cấu hình Tesseract cho tiếng Trung giản thể
    custom_config = r'--oem 3 --psm 6 -l chi_sim'
    
    # Lưu bản sao của hình ảnh gốc để thêm vào tài liệu Word
    original_image = image.copy()
    
    # Tối ưu hóa hình ảnh cho OCR
    ocr_image = image.copy()
    
    # Tăng độ tương phản cho OCR
    enhancer = ImageEnhance.Contrast(ocr_image)
    ocr_image = enhancer.enhance(1.5)
    
    # Chuyển đổi hình ảnh sang tensor PyTorch nếu MPS khả dụng
    if mps_available:
        try:
            import numpy as np
            
            # Chuyển đổi hình ảnh thành mảng numpy
            img_array = np.array(ocr_image)
            
            # Chuyển đổi thành tensor PyTorch và đưa lên GPU
            device = torch.device("mps")
            img_tensor = torch.from_numpy(img_array).to(device)
            
            # Áp dụng một số tiền xử lý đơn giản trên GPU
            if len(img_tensor.shape) == 3:  # Hình ảnh RGB
                img_tensor = img_tensor.float() / 255.0
                img_tensor = torch.clamp((img_tensor - 0.5) * 1.2 + 0.5, 0, 1)
                img_tensor = (img_tensor * 255).byte()
            
            # Chuyển trở lại CPU để sử dụng với Tesseract
            img_array = img_tensor.cpu().numpy()
            
            # Chuyển lại thành định dạng hình ảnh mà Tesseract có thể xử lý
            ocr_image = Image.fromarray(img_array)
            
            logger.info(f"Đã tiền xử lý hình ảnh trang {page_num+1} trên GPU")
        except Exception as e:
            logger.error(f"Lỗi khi xử lý hình ảnh trên GPU: {str(e)}")
    
    # Thực hiện OCR
    text = pytesseract.image_to_string(ocr_image, config=custom_config, lang='chi_sim')
    
    processing_time = time.time() - start_time
    logger.info(f"Đã xử lý trang {page_num+1}/{total_pages} trong {processing_time:.2f} giây")
    
    # Tối ưu hóa hình ảnh gốc để đưa vào tài liệu Word
    optimized_image = optimize_image(original_image, quality=85, max_width=1000)
    
    # Trả về kết quả
    return page_num, text, optimized_image, processing_time

def convert_pdf_to_docx_with_tesseract_parallel(pdf_file, output_folder=None, dpi=300, max_workers=None):
    """
    Chuyển đổi PDF tiếng Trung sang DOCX sử dụng Tesseract OCR với xử lý đa luồng
    
    :param pdf_file: Đường dẫn đến file PDF
    :param output_folder: Thư mục đầu ra (nếu None, sẽ lưu cùng thư mục với file PDF)
    :param dpi: Độ phân giải khi chuyển PDF sang hình ảnh
    :param max_workers: Số luồng tối đa (None = số lõi CPU)
    :return: Đường dẫn đến file DOCX đã tạo
    """
    total_start_time = time.time()
    
    # Tạo tên file đầu ra
    pdf_name = os.path.basename(pdf_file)
    docx_name = os.path.splitext(pdf_name)[0] + '.docx'
    
    if output_folder:
        os.makedirs(output_folder, exist_ok=True)
        docx_file = os.path.join(output_folder, docx_name)
    else:
        docx_file = os.path.splitext(pdf_file)[0] + '.docx'
    
    # Xác định số lượng workers
    if max_workers is None:
        max_workers = os.cpu_count()
    
    logger.info(f"Bắt đầu chuyển đổi PDF sang Word với {max_workers} luồng xử lý")
    logger.info(f"GPU MPS khả dụng: {mps_available}")
    logger.info(f"File đầu vào: {pdf_file}")
    logger.info(f"File đầu ra: {docx_file}")
    
    # Tạo document Word mới
    doc = Document()
    
    # Thiết lập tùy chọn nén hình ảnh trong Word - cách khác không sử dụng xpath
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Tạo cài đặt nén hình ảnh
        settings_xml = f'<w:imageSettings {nsdecls("w")} w:targetScreenSize="hdScreen" w:jpegQuality="90"/>'
        element = parse_xml(settings_xml)
        
        # Thêm vào tài liệu
        doc.settings._element.append(element)
        logger.info("Đã thiết lập nén hình ảnh trong tài liệu Word")
    except Exception as e:
        logger.warning(f"Không thể thiết lập nén hình ảnh: {str(e)}")
    
    # Thiết lập font mặc định cho document là SimSun (phổ biến cho tiếng Trung)
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10)  # Giảm kích thước font mặc định xuống 10pt
    
    # Thiết lập font cho tiếng Trung
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Thiết lập margins của trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        # Thiết lập hướng trang ngang để có nhiều không gian hơn
        section.orientation = 1  # WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    
    # Tạo thư mục tạm để lưu ảnh
    with tempfile.TemporaryDirectory() as temp_dir:
        logger.info(f"Đang chuyển đổi PDF sang hình ảnh với DPI={dpi}...")
        conversion_start = time.time()
        
        # Chuyển đổi PDF sang hình ảnh với độ phân giải phù hợp
        # Giảm DPI xuống để giảm kích thước hình ảnh nhưng vẫn đảm bảo chất lượng
        images = convert_from_path(
            pdf_file, 
            dpi=dpi,
            output_folder=temp_dir,
            fmt="jpeg",  # Sử dụng JPEG thay vì PNG để giảm kích thước
            thread_count=max_workers,
            jpegopt={"quality": 85, "optimize": True}  # Tối ưu hóa JPEG
        )
        
        conversion_time = time.time() - conversion_start
        logger.info(f"Đã chuyển đổi PDF thành {len(images)} hình ảnh trong {conversion_time:.2f} giây")
        
        # Chuẩn bị danh sách tham số cho xử lý đa luồng
        tasks = [(image, i, len(images)) for i, image in enumerate(images)]
        
        # Xử lý OCR đa luồng
        ocr_start = time.time()
        results = []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Bắt đầu xử lý đa luồng
            future_to_page = {executor.submit(process_page, task): task for task in tasks}
            
            # Thu thập kết quả khi hoàn thành
            for future in concurrent.futures.as_completed(future_to_page):
                try:
                    page_num, text, image, processing_time = future.result()
                    results.append((page_num, text, image))
                except Exception as e:
                    task = future_to_page[future]
                    logger.error(f"Lỗi khi xử lý trang {task[1]+1}: {str(e)}")
        
        ocr_time = time.time() - ocr_start
        logger.info(f"Đã hoàn thành OCR cho tất cả các trang trong {ocr_time:.2f} giây")
        
        # Sắp xếp kết quả theo thứ tự trang
        results.sort(key=lambda x: x[0])
        
        # Tạo file Word từ kết quả OCR
        docx_start = time.time()
        
        for page_num, text, image in results:
            # Thêm số trang ở đầu trang
            page_paragraph = doc.add_paragraph()
            page_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            page_run = page_paragraph.add_run(f"Trang {page_num + 1}")
            page_run.font.name = 'SimSun'
            page_run.bold = True
            page_run.font.size = Pt(14)
            
            # Tạo bảng 1x2 để chứa hình ảnh và văn bản với tỉ lệ 2:1
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            
            # Tính toán chiều rộng cột (sử dụng giá trị cụ thể thay vì tính toán phức tạp)
            # python-docx yêu cầu giá trị là số nguyên
            image_width = int(Cm(16.7))  # Cột ảnh (khoảng 15cm)
            text_width = int(Cm(8.3))   # Cột văn bản (khoảng 10cm)
            
            # Thiết lập chiều rộng cột
            table.columns[0].width = image_width
            table.columns[1].width = text_width
            
            # Lấy các ô trong bảng
            image_cell = table.cell(0, 0)
            text_cell = table.cell(0, 1)
            
            # Thêm hình ảnh vào ô bên trái
            image_paragraph = image_cell.paragraphs[0]
            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Tính toán tỷ lệ khung hình để đảm bảo ảnh vừa với cột
            img_width, img_height = image.size
            aspect_ratio = img_width / img_height
            
            # Tính toán kích thước ảnh sao cho vừa với cột và giữ nguyên tỷ lệ
            image_width_cm = 14  # Giới hạn chiều rộng tối đa là 14cm
            image_height_cm = image_width_cm / aspect_ratio
            
            # Giới hạn chiều cao tối đa để tránh ảnh quá cao
            max_height_cm = 18  # Giới hạn chiều cao tối đa là 18cm
            if image_height_cm > max_height_cm:
                image_height_cm = max_height_cm
                image_width_cm = image_height_cm * aspect_ratio
            
            # logger.info(f"Trang {page_num+1}: Kích thước ảnh gốc {img_width}x{img_height}, " +
            #            f"kích thước trên Word: {image_width_cm:.2f}cm x {image_height_cm:.2f}cm")
            
            # Thêm hình ảnh đã được tối ưu hóa
            with io.BytesIO() as image_stream:
                # Lưu hình ảnh với định dạng JPEG để giảm kích thước
                image.save(image_stream, format='JPEG', quality=85, optimize=True)
                image_stream.seek(0)
                image_run = image_paragraph.add_run()
                image_run.add_picture(image_stream, width=Cm(image_width_cm))
            
            # Thêm văn bản vào ô bên phải với font nhỏ hơn
            text_paragraph = text_cell.paragraphs[0]
            
            # Xử lý text để giữ nguyên định dạng
            paragraphs = text.split('\n\n')
            
            for i, para_text in enumerate(paragraphs):
                if i > 0:
                    text_paragraph = text_cell.add_paragraph()
                
                if para_text.strip():
                    # Xử lý các dòng trong đoạn văn
                    lines = para_text.split('\n')
                    for j, line in enumerate(lines):
                        if j > 0:  # Nếu không phải dòng đầu tiên, thêm line break
                            run = text_paragraph.add_run()
                            run.add_break()
                        
                        # Thêm nội dung dòng với font nhỏ hơn
                        if line.strip():
                            run = text_paragraph.add_run(line)
                            run.font.name = 'SimSun'
                            run.font.size = Pt(9)  # Giảm kích thước font xuống 9pt
                            # Thiết lập font cho tiếng Trung
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            
            # Thêm page break sau mỗi trang (trừ trang cuối)
            # if page_num < len(images) - 1:
            #     doc.add_page_break()
        
        # Lưu file DOCX với tùy chọn giảm kích thước
        doc.save(docx_file)
        docx_time = time.time() - docx_start
        logger.info(f"Đã tạo và lưu file Word trong {docx_time:.2f} giây")
    
    total_time = time.time() - total_start_time
    logger.info(f"Tổng thời gian xử lý: {total_time:.2f} giây")
    logger.info(f"Đã chuyển đổi thành công: {pdf_file} -> {docx_file}")
    
    # Hiển thị kích thước file
    file_size = os.path.getsize(docx_file) / (1024 * 1024)  # Chuyển đổi sang MB
    logger.info(f"Kích thước file DOCX: {file_size:.2f} MB")
    
    return docx_file



def pdf2word(pdf_path, output_folder="app/data/uploads"):
    # Sử dụng 2 CPU cores như yêu cầu
    max_workers = 2
    
    # Hiển thị thông tin về GPU
    if mps_available:
        logger.info("MPS (Metal Performance Shaders) khả dụng trên Mac này")
        device = torch.device("mps")
        logger.info(f"Sử dụng thiết bị: {device}")
    else:
        logger.warning("MPS không khả dụng trên Mac này. Sẽ sử dụng CPU.")
    
    # Chuyển đổi PDF sang DOCX với xử lý đa luồng
    output_docx = convert_pdf_to_docx_with_tesseract_parallel(
        pdf_path,
        output_folder=output_folder,
        dpi=400,  # Giảm DPI xuống 400 để cân bằng giữa chất lượng và kích thước file
        max_workers=max_workers
    )
    
    print(f"Đã hoàn thành chuyển đổi. File Word được lưu tại: {output_docx}")

def get_all_files(folder_path):
    try:
        # Trả về danh sách đầy đủ các file (bao gồm đường dẫn)
        files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
        # print(files)
        return [file for file in files if ".DS_Store" not in file]
    except Exception as e:
        print(f"Lỗi: {e}")
        return []
    
pdf_folder = "app/data/downloads/pdf_exam_extract/HSK6"
output_folder = "app/data/uploads/pdf_exam_transform/HSK6"

for pdf_path in get_all_files(pdf_folder):
    pdf2word(pdf_path, output_folder)
